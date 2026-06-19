# -*- coding: utf-8 -*-
"""單一來源 (YAML) -> Word 講義（教師完整版 / 學生挖空版）。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from render import docx_rich, ANS_RED

MAROON = "8C2740"
BTN = "1F6F78"
CJK = "Microsoft JhengHei"


def _set_cjk(run):
    run.font.name = CJK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK)


def _shade(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _para(doc, text="", size=11, bold=False, color=None, before=4, after=4,
          indent=None, fill=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.left_indent = Cm(indent)
    if fill:
        _shade(p, fill)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        _set_cjk(r)
    return p


def _rich(doc, s, mode, size=11, before=3, after=3, indent=None, fill=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.left_indent = Cm(indent)
    if fill:
        _shade(p, fill)
    docx_rich(p, s, mode=mode)
    for r in p.runs:
        if not r.font.size:
            r.font.size = Pt(size)
        _set_cjk(r)
    return p


def _label(doc, text):
    return _para(doc, text, size=11, bold=True, color=MAROON, before=10, after=3)


def _heading(doc, text, level):
    sizes = {0: 18, 1: 15, 2: 13}
    p = _para(doc, text, size=sizes.get(level, 12), bold=True,
              color=MAROON, before=14 if level <= 1 else 8, after=6)
    return p


def _question(doc, q, mode):
    meta = f'［{q["tag"]}］　難易度：{q["level"]}'
    if q.get("core"):
        meta += f'　解題核心：{q["core"]}'
    _para(doc, meta, size=10.5, bold=True, color="6F1F33", before=10, after=2)
    _rich(doc, "題目：" + q["body"], mode, size=11)
    if q.get("table"):
        _docx_table(doc, q["table"])
    for i, o in enumerate(q.get("options", []), 1):
        _rich(doc, f"({i}) " + o, mode, size=11, indent=0.6, before=1, after=1)
    sol = q.get("solution")
    if sol:
        if sol.get("brief"):
            _rich(doc, f'**{sol.get("brief_label","簡答")}：** ' + sol["brief"], mode,
                  size=11, indent=0.4, before=3, after=1, fill="F4F8F7")
        for st in sol.get("steps", []):
            _rich(doc, st, mode, size=11, indent=0.4, before=1, after=1, fill="F4F8F7")


def _docx_table(doc, tbl):
    t = doc.add_table(rows=1, cols=len(tbl["head"]))
    t.style = "Table Grid"
    for c, h in zip(t.rows[0].cells, tbl["head"]):
        docx_rich(c.paragraphs[0], h, mode="teacher")
        for r in c.paragraphs[0].runs:
            r.bold = True
            _set_cjk(r)
    for row in tbl["rows"]:
        cells = t.add_row().cells
        for c, val in zip(cells, row):
            docx_rich(c.paragraphs[0], val, mode="teacher")
            for r in c.paragraphs[0].runs:
                _set_cjk(r)


def _misconceptions(doc, mis, mode):
    _label(doc, "【常見誤解 × → 正解 ○】")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    h = t.rows[0].cells
    for cell, txt in zip(h, ["常見誤解 ✗", "正確理解 ✓"]):
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(MAROON)
        _set_cjk(r)
    for m in mis:
        cells = t.add_row().cells
        docx_rich(cells[0].paragraphs[0], m["wrong"], mode="teacher")
        docx_rich(cells[1].paragraphs[0], m["right"], mode="teacher")
        for cell in cells:
            for r in cell.paragraphs[0].runs:
                _set_cjk(r)


def _kp(doc, kp, mode):
    _heading(doc, f'{kp["num"]}　{kp["title"]}', level=2)
    _rich(doc, "這個考點在學什麼：" + kp["intro"], mode, size=11, after=4)
    _label(doc, "【必背重點與公式】")
    for p in kp["points"]:
        if isinstance(p, dict):
            _rich(doc, "• **" + p["label"] + "：** " + p["lines"][0], mode,
                  size=11, indent=0.4, before=3, after=0)
            for l in p["lines"][1:]:
                _rich(doc, l, mode, size=11, indent=1.0, before=0, after=1)
        elif isinstance(p, (list, tuple)):
            _rich(doc, "• " + p[0], mode, size=11, indent=0.4, before=2, after=1)
            for s in p[1:]:
                _rich(doc, "－ " + s, mode, size=11, indent=0.9, before=0, after=0)
        else:
            _rich(doc, "• " + p, mode, size=11, indent=0.4, before=1, after=1)
    for tb in kp.get("tables", []):
        if tb.get("title"):
            _label(doc, tb["title"])
        t = doc.add_table(rows=1, cols=len(tb["head"]))
        t.style = "Table Grid"
        for cell, htext in zip(t.rows[0].cells, tb["head"]):
            r = cell.paragraphs[0].add_run(htext); r.bold = True
            r.font.color.rgb = RGBColor.from_string(MAROON); _set_cjk(r)
        for row in tb["rows"]:
            cells = t.add_row().cells
            for cell, val in zip(cells, row):
                docx_rich(cell.paragraphs[0], val, mode="teacher")
                for rr in cell.paragraphs[0].runs:
                    _set_cjk(rr)
    if kp.get("misconceptions"):
        _misconceptions(doc, kp["misconceptions"], mode)
    _label(doc, "【歷屆試題】")
    for q in kp.get("questions", []):
        if "subqs" in q:
            _para(doc, f'［{q["tag"]}］　難易度：{q["level"]}', size=10.5, bold=True,
                  color="6F1F33", before=10, after=2)
            _rich(doc, q["body"], mode)
            for sq in q["subqs"]:
                _para(doc, sq["label"], size=11, bold=True, before=4, after=1)
                if sq.get("body"):
                    _rich(doc, sq["body"], mode)
                s = sq.get("solution")
                if s:
                    if s.get("brief"):
                        _rich(doc, "**簡答：** " + s["brief"], mode, indent=0.4, fill="F4F8F7")
                    for st in s.get("steps", []):
                        _rich(doc, st, mode, indent=0.4, fill="F4F8F7")
        else:
            _question(doc, q, mode)
    if kp.get("strategy"):
        _label(doc, "【解題策略】")
        s = kp["strategy"]
        if isinstance(s, (list, tuple)):
            for x in s:
                _rich(doc, "▸ " + x, mode, indent=0.4, fill="FDF3EE", before=1, after=1)
        else:
            _rich(doc, s, mode, indent=0.4, fill="FDF3EE")


def build_docx(unit, mode, out_path):
    doc = Document()
    # 預設字型
    style = doc.styles["Normal"]
    style.font.name = CJK
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)

    ver = "教師完整版" if mode == "teacher" else "學生挖空版"
    _para(doc, f'115 學測數學・{unit["title"]}　複習講義（{ver}）',
          size=18, bold=True, color=MAROON, before=0, after=8)

    # Part 0
    p0 = unit.get("part0")
    if p0:
        _heading(doc, "Part 0　引起動機：穩拿分的入門單元", level=1)
        _label(doc, "近十年出題趨勢（106–115）")
        yrs = p0["trend_table"]["years"]
        cnts = p0["trend_table"]["counts"]
        t = doc.add_table(rows=2, cols=len(yrs) + 2)
        t.style = "Table Grid"
        hdr = ["年度"] + [str(y) for y in yrs] + ["合計"]
        val = ["題數"] + list(cnts) + [p0["trend_table"]["total"]]
        for cell, txt in zip(t.rows[0].cells, hdr):
            r = cell.paragraphs[0].add_run(txt); r.bold = True; _set_cjk(r)
        for cell, txt in zip(t.rows[1].cells, val):
            r = cell.paragraphs[0].add_run(txt); _set_cjk(r)
        _label(doc, "趨勢解讀")
        for n in p0.get("notes", []):
            _rich(doc, "• " + n, mode, indent=0.4, before=1, after=1)
        if p0.get("map"):
            _label(doc, "考點地圖")
            _rich(doc, p0["map"], mode, indent=0.4, fill="FDF3EE")

    # Part 1
    _heading(doc, f'Part 1　建構概念：{unit.get("part1_label","五大考點")}', level=1)
    for kp in unit["kps"]:
        _kp(doc, kp, mode)

    # Part 2
    p2 = unit.get("part2")
    if p2:
        _heading(doc, f'Part 2　喚起行動：模擬實戰', level=1)
        for g in p2["groups"]:
            _label(doc, g["title"])
            for q in g["questions"]:
                _question(doc, q, mode)

    # Part 3
    p3 = unit.get("part3")
    if p3:
        _heading(doc, "Part 3　快速複習：考前速查", level=1)
        _label(doc, "公式一頁速查")
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        for cell, txt in zip(t.rows[0].cells, ["主題", "必記公式 / 結論"]):
            r = cell.paragraphs[0].add_run(txt); r.bold = True
            r.font.color.rgb = RGBColor.from_string(MAROON); _set_cjk(r)
        for row in p3["ref_table"]:
            cells = t.add_row().cells
            r = cells[0].paragraphs[0].add_run(row["k"]); r.bold = True; _set_cjk(r)
            docx_rich(cells[1].paragraphs[0], row["v"], mode="teacher")
            for rr in cells[1].paragraphs[0].runs:
                _set_cjk(rr)
        _label(doc, "常見誤解總清單（考前自我檢查）")
        for i, c in enumerate(p3.get("checklist", []), 1):
            _rich(doc, f"{i}. " + c, mode, indent=0.4, before=1, after=1)

    doc.save(out_path)
    return out_path
