# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FONT = 'Microsoft JhengHei'
MAROON = RGBColor(0x8c, 0x27, 0x40)
RED    = RGBColor(0xc0, 0x39, 0x2b)
GRAY   = RGBColor(0x66, 0x66, 0x66)
BLUE   = RGBColor(0x1f, 0x6f, 0x78)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = sec.bottom_margin = Mm(18)

normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def set_run(r, size=11, bold=False, color=None, italic=False):
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    if color is not None:
        r.font.color.rgb = color

def para(text='', size=11, bold=False, color=None, before=2, after=2, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        set_run(p.add_run(text), size, bold, color, italic)
    return p

def rich(parts, before=2, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for seg in parts:
        txt = seg[0]; opts = seg[1] if len(seg) > 1 else {}
        set_run(p.add_run(txt), opts.get('size', 11), opts.get('bold', False),
                opts.get('color'), opts.get('italic', False))
    return p

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def cell_text(cell, text, bold=False, color=None, size=10.5):
    cell.paragraphs[0].text = ''
    set_run(cell.paragraphs[0].add_run(text), size, bold, color)

BLANK = '＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿'

def blank_line(label):
    rich([(label, {'bold': True}), ('  ' + BLANK, {'color': GRAY})], after=4)

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D9C2C9')
    pbdr.append(bottom); pPr.append(pbdr)

# ---------- Title ----------
para('互動學習講義 ── 單元資料範本', 18, True, MAROON, before=0, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
para('填好這份（或直接把內容貼給我），我就能快速做成同一風格的互動網頁', 10.5, False, GRAY,
     before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------- 填寫說明 callout ----------
t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = True
c = t.cell(0, 0); shade(c, 'FDF3EE')
c.paragraphs[0].text = ''
set_run(c.paragraphs[0].add_run('填寫說明（很重要，請先看）'), 11, True, MAROON)
notes = [
    '1. 這份是「一個單元」用的；每個單元複製一份來填。考點數、題數都不限，不需要的欄位留空或刪除即可。',
    '2. 要「挖空／點按鈕才出現」的答案，請用【 】框起來。例：m 列 n 行的矩陣記為【m×n】階。',
    '3. 「常見誤解→正解」的正解、以及歷屆／練習的簡答與詳解，我會自動做成按鈕（不用特別標）。',
    '4. 數學照平常打就好：次方 A^n、行列式 det(A)、分數 1/2、矩陣寫成 [[a,b],[c,d]]、根號 √3、角度 60°、θ 等都可以，我會轉成漂亮排版。',
    '5. 填完存成 .docx 或 .txt 都行，或直接把文字貼到對話框。',
]
for ntext in notes:
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(ntext), 10.5, False, None)
para('', after=6)

# ---------- 單元基本資料 ----------
para('單元基本資料', 14, True, MAROON, before=6, after=4)
blank_line('單元名稱（例：矩陣、三角函數）：')
blank_line('適用年級 / 範圍：')
blank_line('一句話副標（可留空）：')
hr()

# ---------- Part 1 ----------
para('Part 1　建構概念：各考點', 14, True, MAROON, before=8, after=2)
para('（以下為「一個考點」的欄位，需要幾個考點就複製幾份）', 10, False, GRAY, before=0, after=6, italic=True)

para('考點 ○：（標題）' + '  ' + BLANK, 12, True, BLUE, before=4, after=3)

blank_line('這個考點在學什麼（1～3 句）：')

para('重點與公式（逐條列，挖空答案用【 】框起來）', 11, True, None, before=4, after=2)
rich([('範例：', {'color': GRAY}), ('兩矩陣相等須「【同階】」且「【對應位置元素全相等】」。',
      {'color': GRAY})], after=3)
for i in range(1, 5):
    rich([('%d. ' % i, {'bold': True}), (BLANK, {'color': GRAY})], after=3)

para('常見誤解 × → 正解 ○（可增列）', 11, True, None, before=6, after=2)
mt = doc.add_table(rows=4, cols=2); mt.style = 'Table Grid'
mt.columns[0].width = Mm(82); mt.columns[1].width = Mm(82)
cell_text(mt.cell(0, 0), '常見誤解 ×', True, MAROON); shade(mt.cell(0, 0), 'F3E3E8')
cell_text(mt.cell(0, 1), '正解 ○（會做成按鈕）', True, MAROON); shade(mt.cell(0, 1), 'F3E3E8')
for r in range(1, 4):
    cell_text(mt.cell(r, 0), '')
    cell_text(mt.cell(r, 1), '')

para('歷屆試題（每題一組，可複製多組）', 11, True, None, before=8, after=2)
blank_line('　年度 / 題型（例：108 數A・選填 A）：')
blank_line('　難易度（★☆☆～★★★）：')
blank_line('　解題核心：')
blank_line('　題目：')
blank_line('　選項（選擇題才填，例：(1)… (2)…）：')
blank_line('　簡答（會做成按鈕）：')
blank_line('　解題關鍵 / 詳解（會做成按鈕）：')

para('解題策略（這個考點的整體提醒）', 11, True, None, before=6, after=2)
para('　' + BLANK, 11, False, GRAY, after=2)
para('　' + BLANK, 11, False, GRAY, after=4)
hr()

# ---------- Part 2 ----------
para('Part 2　模擬實戰練習（可選）', 14, True, MAROON, before=8, after=2)
para('（一題一組，可複製多份；題型：單選 / 多選 / 選填 / 非選擇）', 10, False, GRAY, before=0, after=6, italic=True)
blank_line('練習 ○　難易度：　　考點：')
blank_line('　題型（單選 / 多選 / 選填 / 非選）：')
blank_line('　題目：')
blank_line('　選項（選擇題才填）：')
blank_line('　答案（會做成按鈕）：')
blank_line('　詳解（會做成按鈕）：')
hr()

# ---------- Part 3 ----------
para('Part 3　考前速查（可選）', 14, True, MAROON, before=8, after=2)
para('公式速查表（主題 ｜ 必記公式 / 結論）', 11, True, None, before=2, after=2)
qt = doc.add_table(rows=5, cols=2); qt.style = 'Table Grid'
qt.columns[0].width = Mm(55); qt.columns[1].width = Mm(109)
cell_text(qt.cell(0, 0), '主題', True, MAROON); shade(qt.cell(0, 0), 'F3E3E8')
cell_text(qt.cell(0, 1), '必記公式 / 結論', True, MAROON); shade(qt.cell(0, 1), 'F3E3E8')
for r in range(1, 5):
    cell_text(qt.cell(r, 0), ''); cell_text(qt.cell(r, 1), '')

para('常見誤解總清單（考前自我檢查，可增列）', 11, True, None, before=8, after=2)
for i in range(1, 5):
    rich([('%d. ' % i, {'bold': True}), (BLANK, {'color': GRAY})], after=3)

out = '/sessions/focused-jolly-brahmagupta/mnt/outputs/單元資料範本.docx'
doc.save(out)
print('saved', out)
