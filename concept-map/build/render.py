# -*- coding: utf-8 -*-
"""渲染引擎：把「富文字字串」（文字＋數學式＋挖空）轉成 HTML 或 Word。

富文字語法（單一來源共用）：
  - 行內數學： \\( ... \\)        例：\\(A=\\det(A)\\)
  - 顯示數學： \\[ ... \\]
  - 挖空：     【答案】            答案內可再含 \\( ... \\) 數學
                                   例：記為 【\\(m\\times n\\)】 階

三種輸出：
  - html_rich(s)               -> 網頁（挖空＝可點按 span；數學交給 KaTeX）
  - docx_rich(p, s, mode, ...) -> Word 段落（mode='teacher' 紅字答案／'student' 底線挖空）
"""
import re
import html as _html

# 一次抓「挖空」「附註」「粗體」或「行內/顯示數學」
_TOKEN_RE = re.compile(
    r"【(?P<blank>.*?)】"
    r"|\[\[(?P<note_v>.*?)\|\|(?P<note_n>.*?)\]\]"
    r"|\*\*(?P<bold>.+?)\*\*"
    r"|\\\((?P<im>.*?)\\\)"
    r"|\\\[(?P<dm>.*?)\\\]",
    re.S,
)
# 挖空內部只切數學
_MATH_RE = re.compile(r"\\\((?P<im>.*?)\\\)|\\\[(?P<dm>.*?)\\\]", re.S)


def _tok_math(s):
    out, i = [], 0
    for m in _MATH_RE.finditer(s):
        if m.start() > i:
            out.append(("text", s[i:m.start()]))
        if m.group("im") is not None:
            out.append(("math", m.group("im"), False))
        else:
            out.append(("math", m.group("dm"), True))
        i = m.end()
    if i < len(s):
        out.append(("text", s[i:]))
    return out


def tokenize(s):
    """回傳 token 串：('text',s) / ('math',tex,display) / ('blank',[sub-tokens])。"""
    out, i = [], 0
    for m in _TOKEN_RE.finditer(s):
        if m.start() > i:
            out.append(("text", s[i:m.start()]))
        if m.group("blank") is not None:
            out.append(("blank", _tok_math(m.group("blank"))))
        elif m.group("note_v") is not None:
            out.append(("note", _tok_math(m.group("note_v")), _tok_math(m.group("note_n"))))
        elif m.group("bold") is not None:
            out.append(("bold", _tok_math(m.group("bold"))))
        elif m.group("im") is not None:
            out.append(("math", m.group("im"), False))
        else:
            out.append(("math", m.group("dm"), True))
        i = m.end()
    if i < len(s):
        out.append(("text", s[i:]))
    return out


# ---------------- HTML ----------------
def _esc(t):
    # 只轉義純文字段落；數學維持原樣交給 KaTeX
    return _html.escape(t, quote=False)


def _math_html(tex, display):
    # KaTeX 由瀏覽器先解碼實體再讀取，故須轉義 & < >（矩陣欄分隔 & 尤其重要）
    t = _html.escape(tex, quote=False)
    return ("\\[" + t + "\\]") if display else ("\\(" + t + "\\)")


def _html_inner(tokens):
    buf = []
    for tok in tokens:
        if tok[0] == "text":
            buf.append(_esc(tok[1]))
        elif tok[0] == "math":
            buf.append(_math_html(tok[1], tok[2]))
    return "".join(buf)


def html_rich(s):
    """富文字 -> HTML 片段（挖空為可點按 span）。"""
    buf = []
    for tok in tokenize(s):
        if tok[0] == "text":
            buf.append(_esc(tok[1]))
        elif tok[0] == "math":
            buf.append(_math_html(tok[1], tok[2]))
        elif tok[0] == "bold":
            buf.append(f"<b>{_html_inner(tok[1])}</b>")
        elif tok[0] == "note":
            vis = _html_inner(tok[1])
            note = _html_inner(tok[2])
            buf.append(
                '<span class="note" tabindex="0" onclick="tn(event,this)">'
                f'{vis}<span class="np">{note}</span></span>'
            )
        elif tok[0] == "blank":
            inner = _html_inner(tok[1])
            buf.append(
                '<span class="blank" onclick="tb(this)">'
                '<span class="q">？</span>'
                f'<span class="a">{inner}</span></span>'
            )
    return "".join(buf)


def has_blank(s):
    return "【" in s and "】" in s


# ---------------- Word (docx) ----------------
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

ANS_RED = "C0392B"  # 與網頁 --ans 同色
NOTE_GRAY = "6F6F6F"  # 附註灰字


def _add_note_run(paragraph, text):
    from docx.shared import Pt, RGBColor
    run = paragraph.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(NOTE_GRAY)
    return run


def _colorize_omml(omml_el, hex_color):
    """把 OMML 內每個數學 run (m:r) 上色，讓 Word 方程式顯示為指定顏色。"""
    for r in omml_el.iter(qn("m:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        if rpr.find(qn("w:color")) is None:
            c = OxmlElement("w:color")
            c.set(qn("w:val"), hex_color)
            rpr.append(c)
    return omml_el


def _add_math(paragraph, tex, color=None):
    from tex2omml import omml_element
    el = omml_element(tex)
    if color:
        _colorize_omml(el, color)
    paragraph._p.append(el)


def _add_text(paragraph, text, color=None, bold=False):
    if text == "":
        return
    run = paragraph.add_run(text)
    if color:
        run.font.color.rgb = _rgb(color)
    if bold:
        run.bold = True
    return run


def _rgb(hex_color):
    from docx.shared import RGBColor
    return RGBColor.from_string(hex_color)


def _blank_inner_to_docx(paragraph, sub_tokens, color):
    for tok in sub_tokens:
        if tok[0] == "text":
            _add_text(paragraph, tok[1], color=color, bold=True)
        elif tok[0] == "math":
            _add_math(paragraph, tok[1], color=color)


def _add_underline_gap(paragraph, width_chars):
    """學生版挖空：一段加底線的空白。"""
    run = paragraph.add_run(" " * max(4, width_chars))  # en space
    run.underline = True


def docx_rich(paragraph, s, mode="teacher"):
    """把富文字寫進既有的 docx 段落。

    mode='teacher' -> 挖空答案以紅色顯示（含數學式）
    mode='student' -> 挖空以底線空格呈現（答案隱藏）
    """
    for tok in tokenize(s):
        if tok[0] == "text":
            _add_text(paragraph, tok[1])
        elif tok[0] == "math":
            _add_math(paragraph, tok[1])
        elif tok[0] == "bold":
            for sub in tok[1]:
                if sub[0] == "text":
                    _add_text(paragraph, sub[1], bold=True)
                else:
                    _add_math(paragraph, sub[1])
        elif tok[0] == "note":
            # 可見文字照常；附註以灰色小字夾註（Word 版印出來也帶說明）
            for sub in tok[1]:
                if sub[0] == "text":
                    _add_text(paragraph, sub[1])
                else:
                    _add_math(paragraph, sub[1])
            _add_note_run(paragraph, "（註：")
            for sub in tok[2]:
                if sub[0] == "text":
                    _add_note_run(paragraph, sub[1])
                else:
                    _add_math(paragraph, sub[1], color=NOTE_GRAY)
            _add_note_run(paragraph, "）")
        elif tok[0] == "blank":
            if mode == "teacher":
                _blank_inner_to_docx(paragraph, tok[1], ANS_RED)
            else:
                # 以答案的純文字長度估算挖空寬度
                plain = "".join(t[1] for t in tok[1] if t[0] == "text") or "x"
                _add_underline_gap(paragraph, len(plain) + 3)
    return paragraph
